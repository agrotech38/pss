import streamlit as st
from datetime import datetime
import re
import os
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# ---------- Helper: replace text in document ----------
def replace_text_in_paragraph(paragraph, mapping):
    """
    Replace keys in mapping in a paragraph (handles runs).
    mapping: dict of {token: replacement}
    """
    # Simple approach: operate on full paragraph text then rebuild runs
    # But to preserve basic run formatting, we iterate runs and replace text in each run.
    for run in paragraph.runs:
        for key, val in mapping.items():
            if key in run.text:
                run.text = run.text.replace(key, val)

def replace_text_in_table(table, mapping):
    for row in table.rows:
        for cell in row.cells:
            replace_text_in_block(cell, mapping)

def replace_text_in_block(block, mapping):
    # block may be a Document, _Cell, Header, Footer, etc.
    for paragraph in block.paragraphs:
        replace_text_in_paragraph(paragraph, mapping)
    for table in getattr(block, "tables", []):
        replace_text_in_table(table, mapping)

def apply_replacements(doc, mapping):
    # paragraphs & tables in main document
    replace_text_in_block(doc)

    # headers & footers for all sections
    for section in doc.sections:
        header = section.header
        footer = section.footer
        try:
            replace_text_in_block(header)
        except Exception:
            pass
        try:
            replace_text_in_block(footer)
        except Exception:
            pass

# ---------- Helper: find template path ----------
def find_template_for_code(code):
    """
    Return path to template docx for given code.
    Tries several candidate filenames/paths.
    """
    code = (code or "").strip()
    candidates = []
    if code == "001":
        candidates = [
            "MOD PSS.docx",
            "/mnt/data/MOD PSS.docx",
            "templates/MOD PSS.docx"
        ]
    elif code == "002":
        candidates = [
            "FAR PSS.docx",
            "/mnt/data/FAR PSS.docx",
            "templates/FAR PSS.docx"
        ]
    else:
        return None

    for c in candidates:
        if os.path.exists(c):
            return c
    return None

# ---------- Create document from template and mapping ----------
def create_docx_from_template(template_path, mapping):
    """
    Load a docx template, replace placeholders using mapping, and return bytes.
    mapping keys should match tokens exactly as present in the docx (e.g. '{{DD/MM/YYYY}}' or 'DD/MM/YYYY').
    """
    doc = Document(template_path)
    apply_replacements(doc, mapping)

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()

# ---------- Fallback: create docx from scratch (if template missing) ----------
def create_docx_fallback(date_str, salutation1, full_name, designation, company_name, city_state,
                         salutation2, po_id, custom_line, item_details, letterhead_path=None, seal_path=None):
    """
    Fallback creation — similar to previous code — used if template for a given code is not found.
    """
    doc = Document()
    # optional letterhead
    if letterhead_path and os.path.exists(letterhead_path):
        try:
            doc.add_picture(letterhead_path, width=Inches(6.5))
        except Exception:
            pass
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=2)
    try:
        table.columns[0].width = Inches(4.0)
        table.columns[1].width = Inches(2.5)
    except Exception:
        pass
    left_cell = table.cell(0, 0)
    right_cell = table.cell(0, 1)
    p_left = left_cell.paragraphs[0]
    p_left.add_run("Kindly Att.").font.size = Pt(10)
    p_right = right_cell.paragraphs[0]
    p_right.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    p_right.add_run(f"Date: {date_str}").font.size = Pt(10)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(f"{salutation1} {full_name},").bold = True
    doc.add_paragraph(f"({designation})")
    doc.add_paragraph(company_name + ",")
    doc.add_paragraph(city_state)
    doc.add_paragraph()
    doc.add_paragraph(f"Dear {salutation2},")
    doc.add_paragraph()
    doc.add_paragraph(custom_line)
    doc.add_paragraph(f"P.O. ID: {po_id}")
    doc.add_paragraph()
    for item_label, (code, weight) in item_details.items():
        try:
            weight_str = f"{float(weight):.2f}"
        except Exception:
            weight_str = str(weight)
        doc.add_paragraph(f"{item_label}) {code} - {weight_str} MT")
    doc.add_paragraph()
    doc.add_paragraph("Kindly acknowledge receipt of the same.")
    doc.add_paragraph()
    doc.add_paragraph("Yours Faithfully,").runs[0].font.size = Pt(10)
    doc.add_paragraph()
    if seal_path and os.path.exists(seal_path):
        try:
            p_seal = doc.add_paragraph()
            p_seal.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p_seal.add_run().add_picture(seal_path, width=Inches(1.5))
            doc.add_paragraph()
        except Exception:
            pass
    doc.add_paragraph("Authorised Signatory")
    doc.add_paragraph("Aravally Processed Agrotech Pvt Ltd")

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()

# ---------- Streamlit UI ----------
st.set_page_config(page_title="PSS Maker (Template fill)", page_icon="favicon.png")

# hide default menu/footer
hide_st_style = """
<style>
#MainMenu {visibility: hidden;}
footer {visibility: hidden;}
header {visibility: hidden;}
</style>
"""
st.markdown(hide_st_style, unsafe_allow_html=True)

st.title("PSS Maker — Template fill")

# initialize session state for docx
if "docx_bytes" not in st.session_state:
    st.session_state.docx_bytes = None
    st.session_state.filename = None

# prefilled data (as before)
pre_filled_data = {
    "001": {
        "full_name": "Mahendra Tripathi",
        "designation": "Country General Manager & Director",
        "company_name": "Lamberti India Pvt. Ltd.",
        "city_state": "Rajkot, Gujarat",
        "po_id": "PO012",  # default placeholder as requested
        "custom_message": "Sending you Pre-Shipment sample of Guar Gum Powder Modified."
    },
    "002": {
        "full_name": "Mahendra Tripathi",
        "designation": "Country General Manager & Director",
        "company_name": "Lamberti India Pvt. Ltd.",
        "city_state": "Rajkot, Gujarat",
        "po_id": "PO012",
        "custom_message": "Sending you Pre-Shipment sample of FARINA GUAR 200 MESH 5000 T/C."
    }
}

with st.form("form"):
    date_picker = st.date_input("Pick date (calendar)", value=datetime.today())
    default_date_str = date_picker.strftime("%d/%m/%Y")
    date_text = st.text_input("Or edit date (DD/MM/YYYY)", value=default_date_str,
                              help="Enter date in DD/MM/YYYY. If invalid, the calendar selection will be used.")

    salutation1 = st.selectbox("Salutation1", ["Mr.", "Mrs."])
    user_code = st.text_input("Enter Code to auto-fill details (e.g. 001 or 002)")

    if user_code in pre_filled_data:
        data = pre_filled_data[user_code]
        full_name = st.text_input("Full Name", value=data["full_name"])
        designation = st.text_input("Designation", value=data["designation"])
        company_name = st.text_input("Company Name", value=data["company_name"])
        city_state = st.text_input("City, State", value=data["city_state"])
        po_id = st.text_input("P.O. ID (will be placed into template)", value=data.get("po_id", "PO012"))
        custom_line = st.text_input("Pre-Shipment Sample Properties:", value=data.get("custom_message", "Sending you Pre-Shipment sample of"))
    else:
        full_name = st.text_input("Full Name")
        designation = st.text_input("Designation")
        company_name = st.text_input("Company Name")
        city_state = st.text_input("City, State")
        po_id = st.text_input("P.O. ID (will be placed into template)", value="PO012")
        custom_line = st.text_input("Pre-Shipment Sample Properties:", value="Sending you Pre-Shipment sample of")

    salutation2 = st.selectbox("Salutation2", ["Sir", "Ma’am"])
    total_containers = st.number_input("Total Number of Containers", min_value=1, step=1, value=1)
    current_container = st.number_input("Current Container Number", min_value=1, step=1, value=1)

    # We'll capture up to 4 batch placeholders B1..B4 (these will be replaced by item_codes).
    # If your template expects B1..B4 as batch codes, supply them here.
    b1 = st.text_input("B1 (Batch 1) — placeholder will be replaced in template", value="")
    b2 = st.text_input("B2 (Batch 2) — placeholder will be replaced in template", value="")
    b3 = st.text_input("B3 (Batch 3) — placeholder will be replaced in template", value="")
    b4 = st.text_input("B4 (Batch 4) — placeholder will be replaced in template", value="")

    # Also keep previous item details for fallback or for inclusion in mapping
    num_items = st.selectbox("Number of items (for fallback doc generation)", [1,2,3,4,5,6], index=0)
    item_labels = ['A','B','C','D','E','F']
    item_details = {}
    for i in range(num_items):
        code = st.text_input(f"Item {item_labels[i]} Code", key=f"item_code_{i}")
        weight = st.number_input(f"Item {item_labels[i]} Weight (MT)", min_value=0.0, step=0.1, value=4.50, format="%.2f", key=f"item_wt_{i}")
        item_details[item_labels[i]] = (code, weight)

    submitted = st.form_submit_button("Generate DOCX from template")

if submitted:
    # parse/validate date_text -> dd/mm/yyyy; fallback to date_picker
    date_str_final = None
    dt_match = re.match(r"^(\d{1,2})/(\d{1,2})/(\d{4})$", date_text.strip())
    if dt_match:
        d,m,y = dt_match.groups()
        try:
            dt = datetime(int(y), int(m), int(d))
            date_str_final = dt.strftime("%d/%m/%Y")
        except ValueError:
            date_str_final = None
    if not date_str_final:
        date_str_final = date_picker.strftime("%d/%m/%Y")

    # Prepare mapping of placeholders -> replacements
    # Templates (as you showed) use placeholders like: {{DD/MM/YYYY}}, {{PO012}}, {{B1}} etc.
    # We'll prepare several variants of keys to replace: with braces and without.
    # Use values provided by user or defaults requested by you.
    mapping = {}

    # Fill date replacement
    # Replace tokens: '{{DD/MM/YYYY}}', 'DD/MM/YYYY'
    mapping["{{DD/MM/YYYY}}"] = date_str_final
    mapping["DD/MM/YYYY"] = date_str_final

    # P.O. ID replacement (user-supplied or default "PO012")
    po_value = po_id.strip() if po_id and po_id.strip() else "PO012"
    mapping["{{PO012}}"] = po_value
    mapping["PO012"] = po_value
    mapping["{{P.O. ID}}"] = po_value
    mapping["{{P.O. ID:}}"] = po_value  # in case template uses slightly different

    # Batch placeholders B1..B4
    b_vals = [b1.strip(), b2.strip(), b3.strip(), b4.strip()]
    # If user didn't fill B1..B4, try to use provided item codes as fallback
    fallback_codes = [item_details.get(lbl, ("",0))[0] for lbl in item_labels]
    for i in range(4):
        val = b_vals[i] if b_vals[i] else (fallback_codes[i] if i < len(fallback_codes) else "")
        token_braced = f"{{{{B{i+1}}}}}"   # e.g. '{{B1}}'
        token_plain = f"B{i+1}"            # e.g. 'B1'
        mapping[token_braced] = val
        mapping[token_plain] = val

    # Also add replacements for placeholders shown in your screenshots such as '{{{B1}}}' etc.
    for i in range(4):
        triple = f"{{{{{{B{i+1}}}}}}}"  # '{{{B1}}}' if present
        if triple not in mapping:
            mapping[triple] = mapping.get(f"{{{{B{i+1}}}}}", "")

    # Find template for the entered code
    template_path = find_template_for_code(user_code)

    # Provide candidate seal/letterhead paths (optional)
    letterhead_path = "letterhead.png" if os.path.exists("letterhead.png") else None
    seal_candidates = ["/mnt/data/APAPL SEAL.png", "APAPL SEAL.png", "APAPL_SEAL.png", "seal.png"]
    seal_path = None
    for cand in seal_candidates:
        if os.path.exists(cand):
            seal_path = cand
            break

    # Compute output filename
    suffix = "MOD" if user_code == "001" else "FAR" if user_code == "002" else "GEN"
    safe_po = re.sub(r'[\/:*?"<>|]', '', po_value)
    po_suffix = safe_po[-3:] if len(safe_po) >= 3 else "000"
    filename = f"PSS LIPL {suffix} {po_suffix} {int(current_container)} of {int(total_containers)}.docx"

    # If template exists, use it and apply replacements
    final_bytes = None
    if template_path:
        try:
            # Prepare mapping keys to include common variants found in doc (with/without braces)
            # (mapping already contains several common forms)
            final_bytes = create_docx_from_template(template_path, mapping)
            st.success(f"Template used: {template_path}")
        except Exception as e:
            st.error(f"Error applying template replacements: {e}")
            # fallback to generating from scratch
            final_bytes = create_docx_fallback(date_str_final, salutation1, full_name, designation, company_name, city_state,
                                               salutation2, po_value, custom_line, item_details,
                                               letterhead_path=letterhead_path, seal_path=seal_path)
    else:
        st.warning("Template for this code not found — generating fallback DOCX.")
        final_bytes = create_docx_fallback(date_str_final, salutation1, full_name, designation, company_name, city_state,
                                           salutation2, po_value, custom_line, item_details,
                                           letterhead_path=letterhead_path, seal_path=seal_path)

    # store in session and present download
    st.session_state.docx_bytes = final_bytes
    st.session_state.filename = filename

# Always show download button if docx_bytes exists
if st.session_state.get("docx_bytes"):
    st.download_button(
        "Download filled DOCX",
        st.session_state.docx_bytes,
        file_name=st.session_state.filename,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
